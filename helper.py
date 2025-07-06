# # helper.py
# import fitz  # PyMuPDF
# import docx

# def extract_text_from_pdf(file):
#     text = ""
#     pdf = fitz.open(stream=file.read(), filetype="pdf")
#     for page in pdf:
#         text += page.get_text()
#     return text

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith(".docx"):
#         return extract_text_from_docx(file)
#     elif file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     else:
#         return "Unsupported file format"

# import fitz  # PyMuPDF
# import docx
# from docx import Document
# import datetime
# import json
# import spacy
# import spacy
# nlp = spacy.load("en_core_web_sm")



# def extract_text_from_pdf(file):
#     text = ""
#     pdf = fitz.open(stream=file.read(), filetype="pdf")
#     for page in pdf:
#         text += page.get_text()
#     return text

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text(file):
#     if file.name.endswith(".pdf"):
#         return extract_text_from_pdf(file)
#     elif file.name.endswith(".docx"):
#         return extract_text_from_docx(file)
#     elif file.name.endswith(".txt"):
#         return file.read().decode("utf-8")
#     else:
#         return "Unsupported file format"

# # ✅ Export summary & Q&A to Word
# def export_to_docx(summary: str, answers: list[str], filename="legal_summary.docx"):
#     doc = Document()
    
#     doc.add_heading("Legal Document Summary", level=1)
#     doc.add_paragraph(summary)

#     if answers:
#         doc.add_heading("Q&A", level=1)
#         for i, answer in enumerate(answers, start=1):
#             doc.add_paragraph(f"{i}. {answer}")

#     doc.add_paragraph(f"\nExported on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
#     doc.save(filename)
#     return filename


# def highlight_legal_terms(text, glossary_file='glossary.json'):
#     with open(glossary_file, 'r') as f:
#         glossary = json.load(f)
#     for term, definition in glossary.items():
#         text = text.replace(
#             term,
#             f"<span title='{definition}' style='background-color:#fce4ec;'>{term}</span>"
#         )
#     return text

# def highlight_named_entities(text):
#     doc = nlp(text)
#     highlighted = text
#     for ent in doc.ents:
#         highlighted = highlighted.replace(ent.text, f"<mark style='background-color:#e0f7fa;'>{ent.text}</mark>")
#     return highlighted


# from docx import Document
# import textract
# import fitz  # PyMuPDF
# import spacy

# # Load spaCy model once
# nlp = spacy.load("en_core_web_sm")

# # Extract text from uploaded file
# def extract_text(uploaded_file):
#     file_type = uploaded_file.name.split('.')[-1]

#     if file_type == "pdf":
#         text = ""
#         doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#         for page in doc:
#             text += page.get_text()
#         return text

#     elif file_type == "docx":
#         return textract.process(uploaded_file).decode("utf-8")

#     elif file_type == "txt":
#         return uploaded_file.read().decode("utf-8")

#     else:
#         return "Unsupported file format."

# # Export summary and Q&A to Word doc
# def export_to_docx(summary, qa_list):
#     doc = Document()
#     doc.add_heading("Legal Document Summary", level=1)
#     doc.add_paragraph(summary)

#     if qa_list:
#         doc.add_heading("Q&A", level=1)
#         for q, a in qa_list:
#             doc.add_paragraph(f"Q: {q}", style='List Bullet')
#             doc.add_paragraph(f"A: {a}", style='List Bullet')

#     doc.save("Legal_Summary.docx")

# # ✅ NER function — fix the import error
# def extract_named_entities(text):
#     doc = nlp(text)
#     return [(ent.text, ent.label_) for ent in doc.ents]



from docx import Document
import textract
import fitz  # PyMuPDF
# import spacy
from PyPDF2 import PdfReader
import docx


# from spacy.cli import download
# # Load spaCy model, downloading if necessary
# try:
#     nlp = spacy.load("en_core_web_sm")
# except OSError:
#     download("en_core_web_sm")
#     nlp = spacy.load("en_core_web_sm")



# Extract text from uploaded file
def extract_text(uploaded_file):
    file_type = uploaded_file.name.split('.')[-1]

    if file_type == "pdf":
        text = ""
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        for page in doc:
            text += page.get_text()
        return text

    elif file_type == "docx":
        return textract.process(uploaded_file).decode("utf-8")

    elif file_type == "txt":
        return uploaded_file.read().decode("utf-8")

    else:
        return "Unsupported file format."

# Export summary and Q&A to Word
def export_to_docx(summary, qa_list):
    doc = Document()
    doc.add_heading("Legal Document Summary", level=1)
    doc.add_paragraph(summary)

    if qa_list:
        doc.add_heading("Q&A", level=1)
        for question, answer in qa_list:
            doc.add_paragraph(f"Q: {question}", style='List Bullet')
            doc.add_paragraph(f"A: {answer}", style='List Bullet')

    doc.save("Legal_Summary.docx")

# Named Entity Recognition
# def extract_named_entities(text):
#     doc = nlp(text)
#     return [(ent.text, ent.label_) for ent in doc.ents]



def check_document_authenticity(file):
    file_type = file.name.split('.')[-1]
    meta_info = {}

    if file_type == "pdf":
        reader = PdfReader(file)
        metadata = reader.metadata
        if metadata:
            meta_info = {key[1:]: str(value) for key, value in metadata.items()}
    
    elif file_type == "docx":
        document = docx.Document(file)
        props = document.core_properties
        meta_info = {
            "author": props.author,
            "title": props.title,
            "created": props.created,
            "last_modified_by": props.last_modified_by
        }

    return meta_info